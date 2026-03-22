import io
from pathlib import Path


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file using PyMuPDF (fitz).
    Handles scanned PDFs gracefully.
    """
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text = []

        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages_text.append(f"[Page {page_num + 1}]\n{text.strip()}")

        doc.close()

        if not pages_text:
            return ""

        return "\n\n".join(pages_text)

    except ImportError:
        raise ImportError(
            "PyMuPDF not installed. Run: pip install PyMuPDF"
        )
    except Exception as e:
        raise ValueError(f"Could not read PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file using python-docx.
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)

    except ImportError:
        raise ImportError(
            "python-docx not installed. Run: pip install python-docx"
        )
    except Exception as e:
        raise ValueError(f"Could not read DOCX: {str(e)}")


def extract_text(uploaded_file) -> tuple[str, str]:
    """
    Auto-detect file type and extract text.

    Returns:
        tuple: (extracted_text, file_type)
    """
    file_bytes = uploaded_file.read()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
        return text, "PDF"

    elif filename.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
        return text, "DOCX"

    elif filename.endswith(".txt"):
        # Try UTF-8 first, then fallback for Devanagari
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("utf-8-sig", errors="replace")
        return text, "TXT"

    else:
        raise ValueError(f"Unsupported file type: {filename}")


def get_word_count(text: str) -> int:
    """Count words in text (works for both Latin and Devanagari)."""
    if not text:
        return 0
    # Split on whitespace for both scripts
    words = text.split()
    return len(words)


def truncate_for_api(text: str, max_chars: int = 7000) -> str:
    """
    Truncate text to fit within API limits.
    Tries to cut at a paragraph boundary.
    """
    if len(text) <= max_chars:
        return text

    # Try to cut at paragraph boundary
    truncated = text[:max_chars]
    last_newline = truncated.rfind("\n\n")
    if last_newline > max_chars * 0.8:
        return truncated[:last_newline] + "\n\n[... document truncated for analysis ...]"

    return truncated + "\n\n[... document truncated for analysis ...]"
