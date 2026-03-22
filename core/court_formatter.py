"""
core/court_formatter.py

Generates properly formatted .docx files for Indian court documents.
Each court type has its own formatting template:
  - Header:  Court name, case number, party names
  - Body:    Times New Roman 12pt, double spacing, justified
  - Footer:  Page numbers, advocate details
  - Margins: As per court rules (typically 1.5" left, 1" others)

Supported formats:
  - Supreme Court of India
  - High Courts (Bombay, Delhi, Allahabad, Madras, Calcutta, etc.)
  - District / Sessions Courts
  - Specialised Tribunals
  - General legal documents (NDA, contracts, affidavits, notices)
"""

import io
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Court format profiles ─────────────────────────────────────────────────────
COURT_PROFILES = {
    "Supreme Court of India": {
        "margin_left": 1.75, "margin_right": 1.0,
        "margin_top": 1.25, "margin_bottom": 1.0,
        "font_body": "Times New Roman", "font_heading": "Times New Roman",
        "size_body": 12, "size_heading": 14,
        "line_spacing": 1.5,
        "para_spacing_before": 0, "para_spacing_after": 6,
        "header_border": True, "footer_border": True,
        "case_prefix": "In the Supreme Court of India",
        "salutation": "MOST RESPECTFULLY SHOWETH:",
        "prayer_header": "PRAYER",
        "verification_required": True,
        "justify": True,
    },
    "Bombay High Court": {
        "margin_left": 1.75, "margin_right": 1.0,
        "margin_top": 1.25, "margin_bottom": 1.0,
        "font_body": "Times New Roman", "font_heading": "Times New Roman",
        "size_body": 12, "size_heading": 14,
        "line_spacing": 1.5,
        "para_spacing_before": 0, "para_spacing_after": 6,
        "header_border": True, "footer_border": True,
        "case_prefix": "In the High Court of Judicature at Bombay",
        "salutation": "MOST RESPECTFULLY SHOWETH:",
        "prayer_header": "PRAYER",
        "verification_required": True,
        "justify": True,
    },
    "Delhi High Court": {
        "margin_left": 1.75, "margin_right": 1.0,
        "margin_top": 1.25, "margin_bottom": 1.0,
        "font_body": "Times New Roman", "font_heading": "Times New Roman",
        "size_body": 12, "size_heading": 14,
        "line_spacing": 1.5,
        "para_spacing_before": 0, "para_spacing_after": 6,
        "header_border": True, "footer_border": True,
        "case_prefix": "In the High Court of Delhi at New Delhi",
        "salutation": "MOST RESPECTFULLY SHOWETH:",
        "prayer_header": "PRAYER",
        "verification_required": True,
        "justify": True,
    },
    "District Court": {
        "margin_left": 1.5, "margin_right": 1.0,
        "margin_top": 1.0, "margin_bottom": 1.0,
        "font_body": "Times New Roman", "font_heading": "Times New Roman",
        "size_body": 12, "size_heading": 13,
        "line_spacing": 1.5,
        "para_spacing_before": 0, "para_spacing_after": 6,
        "header_border": False, "footer_border": False,
        "case_prefix": "In the Court of",
        "salutation": "RESPECTFULLY SHOWETH:",
        "prayer_header": "PRAYER",
        "verification_required": True,
        "justify": True,
    },
    "Family Court": {
        "margin_left": 1.5, "margin_right": 1.0,
        "margin_top": 1.0, "margin_bottom": 1.0,
        "font_body": "Times New Roman", "font_heading": "Times New Roman",
        "size_body": 12, "size_heading": 13,
        "line_spacing": 1.5,
        "para_spacing_before": 0, "para_spacing_after": 6,
        "header_border": False, "footer_border": False,
        "case_prefix": "In the Family Court",
        "salutation": "RESPECTFULLY SHOWETH:",
        "prayer_header": "PRAYER",
        "verification_required": True,
        "justify": True,
    },
    "NCLT": {
        "margin_left": 1.5, "margin_right": 1.0,
        "margin_top": 1.0, "margin_bottom": 1.0,
        "font_body": "Arial", "font_heading": "Arial",
        "size_body": 11, "size_heading": 13,
        "line_spacing": 1.5,
        "para_spacing_before": 0, "para_spacing_after": 6,
        "header_border": False, "footer_border": False,
        "case_prefix": "Before the National Company Law Tribunal",
        "salutation": "RESPECTFULLY SHOWETH:",
        "prayer_header": "PRAYER",
        "verification_required": True,
        "justify": True,
    },
    "General Legal Document": {
        "margin_left": 1.25, "margin_right": 1.25,
        "margin_top": 1.0, "margin_bottom": 1.0,
        "font_body": "Times New Roman", "font_heading": "Times New Roman",
        "size_body": 12, "size_heading": 14,
        "line_spacing": 1.5,
        "para_spacing_before": 0, "para_spacing_after": 6,
        "header_border": False, "footer_border": False,
        "case_prefix": "",
        "salutation": "",
        "prayer_header": "CONCLUSION",
        "verification_required": False,
        "justify": True,
    },
}

# Fallback profile for unknown courts
DEFAULT_PROFILE = COURT_PROFILES["Bombay High Court"]


def _get_profile(court: str) -> dict:
    """Match court name to closest profile."""
    if not court:
        return DEFAULT_PROFILE
    court_lower = court.lower()
    if "supreme" in court_lower:
        return COURT_PROFILES["Supreme Court of India"]
    if "bombay" in court_lower or "mumbai" in court_lower:
        return COURT_PROFILES["Bombay High Court"]
    if "delhi" in court_lower:
        return COURT_PROFILES["Delhi High Court"]
    if "high court" in court_lower or "high c" in court_lower:
        # Generic HC — use Bombay profile as standard
        return COURT_PROFILES["Bombay High Court"]
    if "family" in court_lower:
        return COURT_PROFILES["Family Court"]
    if "nclt" in court_lower or "company law" in court_lower:
        return COURT_PROFILES["NCLT"]
    if "district" in court_lower or "sessions" in court_lower or "civil court" in court_lower:
        return COURT_PROFILES["District Court"]
    if "tribunal" in court_lower or "forum" in court_lower or "drt" in court_lower:
        return COURT_PROFILES["District Court"]
    # NDA, contracts, notices — no court
    return COURT_PROFILES["General Legal Document"]


def _set_page_margins(section, profile: dict):
    """Apply court-specific page margins."""
    section.left_margin   = Inches(profile["margin_left"])
    section.right_margin  = Inches(profile["margin_right"])
    section.top_margin    = Inches(profile["margin_top"])
    section.bottom_margin = Inches(profile["margin_bottom"])


def _set_run_font(run, font_name: str, size_pt: int,
                  bold: bool = False, italic: bool = False, underline: bool = False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.underline = underline


def _add_paragraph(doc, text: str, profile: dict,
                   bold: bool = False, italic: bool = False,
                   align: str = "justify", size_override: int = None,
                   space_before: int = None, space_after: int = None,
                   keep_together: bool = False) -> object:
    """Add a formatted paragraph with court-standard styling."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before if space_before is not None else profile["para_spacing_before"])
    para.paragraph_format.space_after  = Pt(space_after  if space_after  is not None else profile["para_spacing_after"])
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = profile["line_spacing"]
    if keep_together:
        para.paragraph_format.keep_together = True

    align_map = {
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = align_map.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)

    if text:
        run = para.add_run(text)
        _set_run_font(run, profile["font_body"],
                      size_override or profile["size_body"], bold, italic)
    return para


def _add_heading(doc, text: str, profile: dict, level: int = 1,
                 underline: bool = False, align: str = "center") -> object:
    """Add a bold, centred heading."""
    para = _add_paragraph(doc, text, profile, bold=True,
                          align=align,
                          size_override=profile["size_heading"],
                          space_before=6, space_after=6)
    if underline and para.runs:
        para.runs[0].underline = True
    return para


def _add_horizontal_rule(doc):
    """Add a thin horizontal line (paragraph bottom border)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _add_page_numbers(section, profile: dict, court: str,
                       doc_type: str, advocate_name: str):
    """Add header and footer with page numbers."""
    header = section.header
    footer = section.footer

    # ── Header ────────────────────────────────────────────────────────────────
    if header.paragraphs:
        hp = header.paragraphs[0]
    else:
        hp = header.add_paragraph()

    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run(doc_type or "Legal Document")
    run.font.name = profile["font_heading"]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    if profile.get("header_border"):
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "000000")
        pBdr.append(bottom)
        pPr.append(pBdr)

    # ── Footer ────────────────────────────────────────────────────────────────
    if footer.paragraphs:
        fp = footer.paragraphs[0]
    else:
        fp = footer.add_paragraph()

    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Left: advocate name
    if advocate_name:
        left_run = fp.add_run(f"{advocate_name}  ")
        left_run.font.name = profile["font_body"]
        left_run.font.size = Pt(9)

    # Centre: page number field
    fld_run = fp.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    fld_run._r.append(fldChar1)
    fld_run._r.append(instrText)
    fld_run._r.append(fldChar2)
    fld_run.font.name = profile["font_body"]
    fld_run.font.size = Pt(9)

    # Separator + total
    sep_run = fp.add_run("  /  ")
    sep_run.font.name = profile["font_body"]
    sep_run.font.size = Pt(9)

    fld_run2 = fp.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.text = " NUMPAGES "
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    fld_run2._r.append(fldChar3)
    fld_run2._r.append(instrText2)
    fld_run2._r.append(fldChar4)
    fld_run2.font.name = profile["font_body"]
    fld_run2.font.size = Pt(9)

    if profile.get("footer_border"):
        pPr = fp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "4")
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), "000000")
        pBdr.append(top)
        pPr.append(pBdr)


def _parse_draft_into_sections(draft_text: str) -> dict:
    """
    Parse Claude's drafted document text into logical sections.
    Returns dict with: cause_title, body_paragraphs, prayer, verification
    """
    sections = {
        "cause_title_lines": [],
        "body_paragraphs":   [],
        "prayer_lines":      [],
        "verification_lines":[],
        "annexure_lines":    [],
    }

    lines = draft_text.strip().split("\n")
    current_section = "cause_title"

    prayer_markers    = {"prayer", "wherefore", "it is therefore prayed", "it is humbly prayed", "प्रार्थना", "विनंती"}
    verify_markers    = {"verification", "verified", "सत्यापन", "पुष्टीकरण", "शपथपत्र"}
    annexure_markers  = {"annexure", "exhibit", "schedule", "list of annexures", "अनुलग्नक", "परिशिष्ट"}
    body_start_markers= {"respectfully showeth", "most respectfully showeth", "the petitioner", "the applicant",
                         "the plaintiff", "the complainant", "1.", "background", "facts"}

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            if current_section == "body":
                sections["body_paragraphs"].append("")
            continue

        ll = line.lower()

        # Detect section transitions
        if any(m in ll for m in prayer_markers):
            current_section = "prayer"
        elif any(m in ll for m in verify_markers):
            current_section = "verification"
        elif any(m in ll for m in annexure_markers):
            current_section = "annexure"
        elif current_section == "cause_title" and any(m in ll for m in body_start_markers):
            current_section = "body"

        # Append to appropriate bucket
        if current_section == "cause_title":
            sections["cause_title_lines"].append(line)
        elif current_section == "body":
            sections["body_paragraphs"].append(line)
        elif current_section == "prayer":
            sections["prayer_lines"].append(line)
        elif current_section == "verification":
            sections["verification_lines"].append(line)
        elif current_section == "annexure":
            sections["annexure_lines"].append(line)

    return sections


def _is_numbered_paragraph(line: str) -> bool:
    """Detect numbered paragraphs like '1.', '(a)', '(i)', etc."""
    return bool(re.match(r"^\s*(\d+\.|[a-z]\.|[ivxIVX]+\.|[\(\[][\w]+[\)\]])", line.strip()))


def _indent_for_numbered(para, level: int = 1):
    """Set paragraph indent for numbered items."""
    para.paragraph_format.left_indent  = Inches(level * 0.5)
    para.paragraph_format.first_line_indent = Inches(-0.35)


# ── Main public function ──────────────────────────────────────────────────────

def create_court_document(
    draft_text: str,
    court: str = "",
    doc_type: str = "",
    party_details: str = "",
    case_number: str = "",
    state: str = "",
    advocate_name: str = "",
    advocate_enrolment: str = "",
    language: str = "English",
) -> bytes:
    """
    Convert draft text into a properly formatted Indian court .docx file.

    Returns bytes of the .docx file ready for st.download_button.
    """
    profile = _get_profile(court)
    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    _set_page_margins(section, profile)
    section.page_width  = Cm(21.0)   # A4 width
    section.page_height = Cm(29.7)   # A4 height

    # ── Header / Footer ────────────────────────────────────────────────────────
    _add_page_numbers(section, profile, court, doc_type, advocate_name)

    # ── Parse the draft ────────────────────────────────────────────────────────
    parsed = _parse_draft_into_sections(draft_text)

    # ── CAUSE TITLE ────────────────────────────────────────────────────────────
    if profile.get("case_prefix"):
        _add_heading(doc, profile["case_prefix"].upper(), profile,
                     align="center", underline=False)

    # Court name
    court_display = court.upper() if court else "COURT OF LAW"
    _add_heading(doc, court_display, profile, align="center", underline=True)
    if state:
        _add_paragraph(doc, f"At {state}", profile, align="center",
                       size_override=profile["size_body"])

    # Case type + number
    if doc_type or case_number:
        case_line = f"{doc_type.upper()}" if doc_type else ""
        if case_number:
            case_line += f"  NO. {case_number}" if case_line else f"NO. {case_number}"
        _add_paragraph(doc, case_line, profile, bold=True, align="center",
                       space_before=8, space_after=4)

    _add_paragraph(doc, "IN THE MATTER OF:", profile, bold=True,
                   align="center", space_before=8)
    _add_horizontal_rule(doc)

    # Parties from cause_title section or party_details
    cause_lines = parsed["cause_title_lines"]
    if cause_lines:
        # Use parsed lines — skip lines that were already printed above
        for line in cause_lines:
            ll = line.lower()
            if any(x in ll for x in ["high court", "supreme court", "district court",
                                      "in the", "at bombay", "at new delhi", "at mumbai"]):
                continue
            _add_paragraph(doc, line, profile, align="justify")
    elif party_details:
        # Format "ABC vs XYZ" party string
        if " vs " in party_details.lower():
            parts = re.split(r"\s+vs\.?\s+", party_details, flags=re.IGNORECASE)
            if len(parts) >= 2:
                _add_paragraph(doc, parts[0].strip(), profile, align="justify",
                               bold=True, space_before=4)
                _add_paragraph(doc, "... PETITIONER / APPELLANT", profile,
                               align="right", italic=True)
                _add_paragraph(doc, "VERSUS", profile, align="center",
                               bold=True, space_before=4, space_after=4)
                _add_paragraph(doc, parts[1].strip(), profile, align="justify", bold=True)
                _add_paragraph(doc, "... RESPONDENT", profile,
                               align="right", italic=True)

    _add_horizontal_rule(doc)

    # Document title
    if doc_type:
        _add_heading(doc, doc_type.upper(), profile,
                     align="center", underline=True)
        _add_paragraph(doc, "", profile)  # spacer

    # ── BODY ───────────────────────────────────────────────────────────────────
    if profile.get("salutation") and parsed["body_paragraphs"]:
        _add_paragraph(doc, profile["salutation"], profile, bold=True,
                       align="left", space_before=8)

    for line in parsed["body_paragraphs"]:
        if not line:
            _add_paragraph(doc, "", profile, space_before=0, space_after=2)
            continue
        is_numbered = _is_numbered_paragraph(line)
        para = _add_paragraph(doc, line, profile,
                              align="justify",
                              bold=False)
        if is_numbered:
            _indent_for_numbered(para)

    # If body is empty (document is all one block), render full draft as body
    if not parsed["body_paragraphs"] and not parsed["cause_title_lines"]:
        for line in draft_text.strip().split("\n"):
            if line.strip():
                _add_paragraph(doc, line.strip(), profile, align="justify")
            else:
                _add_paragraph(doc, "", profile, space_before=0, space_after=2)

    # ── PRAYER ─────────────────────────────────────────────────────────────────
    if parsed["prayer_lines"]:
        _add_paragraph(doc, "", profile)
        _add_heading(doc, profile["prayer_header"], profile,
                     align="center", underline=True)
        for line in parsed["prayer_lines"]:
            if not line:
                continue
            ll = line.lower()
            if ll in {"prayer", "wherefore", "प्रार्थना", "विनंती"}:
                continue  # already printed as heading
            is_sub = _is_numbered_paragraph(line)
            para = _add_paragraph(doc, line, profile, align="justify")
            if is_sub:
                _indent_for_numbered(para)

    # ── SIGNATURE BLOCK ────────────────────────────────────────────────────────
    _add_paragraph(doc, "", profile, space_before=12)
    sig_para = _add_paragraph(doc, "Place: _________________________", profile,
                              align="left", space_before=4, space_after=2)
    _add_paragraph(doc, "Date:  _________________________", profile,
                  align="left", space_before=2, space_after=12)

    # Advocate signature block (right-aligned)
    sig_right = doc.add_paragraph()
    sig_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_right.paragraph_format.space_before = Pt(0)
    sig_right.paragraph_format.space_after  = Pt(2)
    r = sig_right.add_run("_" * 35)
    r.font.name = profile["font_body"]
    r.font.size = Pt(profile["size_body"])

    adv_para = doc.add_paragraph()
    adv_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    adv_para.paragraph_format.space_before = Pt(2)
    adv_para.paragraph_format.space_after  = Pt(0)
    adv_run = adv_para.add_run(advocate_name if advocate_name else "Advocate for the Petitioner")
    adv_run.font.name = profile["font_body"]
    adv_run.font.size = Pt(profile["size_body"])
    adv_run.bold = True

    if advocate_enrolment:
        enrol_para = doc.add_paragraph()
        enrol_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        enrol_para.paragraph_format.space_before = Pt(0)
        enrol_para.paragraph_format.space_after  = Pt(0)
        er = enrol_para.add_run(f"Enrolment No.: {advocate_enrolment}")
        er.font.name = profile["font_body"]
        er.font.size = Pt(9)

    # ── VERIFICATION ───────────────────────────────────────────────────────────
    if profile.get("verification_required") and parsed["verification_lines"]:
        doc.add_page_break()
        _add_heading(doc, "VERIFICATION", profile, align="center", underline=True)
        for line in parsed["verification_lines"]:
            if line and line.lower() not in {"verification", "verified"}:
                _add_paragraph(doc, line, profile, align="justify")
        _add_paragraph(doc, "", profile, space_before=16)
        _add_paragraph(doc, "Verified at __________ on this _____ day of __________, 20__",
                       profile, align="left")
        _add_paragraph(doc, "", profile, space_before=12)
        verif_sig = doc.add_paragraph()
        verif_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        verif_sig.paragraph_format.space_before = Pt(0)
        vr = verif_sig.add_run("Deponent / Petitioner")
        vr.font.name = profile["font_body"]
        vr.font.size = Pt(profile["size_body"])
        vr.bold = True

    elif profile.get("verification_required"):
        # Add a blank verification template
        doc.add_page_break()
        _add_heading(doc, "VERIFICATION", profile, align="center", underline=True)
        _add_paragraph(
            doc,
            "I, the above-named Petitioner / Deponent, do hereby solemnly affirm and "
            "verify that the contents of the above paragraphs are true to my knowledge "
            "and belief and that nothing material has been concealed therefrom.",
            profile, align="justify"
        )
        _add_paragraph(doc, "", profile, space_before=16)
        _add_paragraph(doc, "Verified at __________ on this _____ day of __________, 20__",
                       profile, align="left")
        _add_paragraph(doc, "", profile, space_before=12)
        verif_sig2 = doc.add_paragraph()
        verif_sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vr2 = verif_sig2.add_run("Deponent / Petitioner")
        vr2.font.name = profile["font_body"]
        vr2.font.size = Pt(profile["size_body"])
        vr2.bold = True

    # ── ANNEXURE LIST ──────────────────────────────────────────────────────────
    if parsed["annexure_lines"]:
        doc.add_page_break()
        _add_heading(doc, "LIST OF ANNEXURES", profile, align="center", underline=True)
        for line in parsed["annexure_lines"]:
            if line.lower() not in {"list of annexures", "annexure", "exhibits"}:
                _add_paragraph(doc, line, profile, align="left")

    # ── Save to bytes ──────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def get_court_profile_summary(court: str) -> dict:
    """Return human-readable format summary for display in the UI."""
    profile = _get_profile(court)
    return {
        "font":         f"{profile['font_body']} {profile['size_body']}pt",
        "line_spacing": f"{profile['line_spacing']}x",
        "margins":      f"Left {profile['margin_left']}\" · Right {profile['margin_right']}\" · Top {profile['margin_top']}\"",
        "header":       "Yes — court name + document type" if profile.get("header_border") else "Simple",
        "footer":       "Page X / Y",
        "verification": "Included" if profile.get("verification_required") else "Not applicable",
    }
